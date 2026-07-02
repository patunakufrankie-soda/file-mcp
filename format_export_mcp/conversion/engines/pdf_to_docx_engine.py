from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from .base import BaseEngine, ConversionResult

logger = logging.getLogger(__name__)


def _detect_table_regions(page) -> list[tuple[float, float, float, float]]:
    """Detect table regions on a page using heuristics.

    Returns list of bounding boxes (x0, y0, x1, y1) for detected tables.
    """
    tables = []

    # Get all text blocks with positions
    blocks = (
        page.get_text("dict").get("blocks", [])
        if isinstance(page.get_text("dict"), dict)
        else []
    )

    # Look for grid-like patterns: multiple blocks aligned horizontally and vertically
    # Simple heuristic: find groups of blocks with similar Y coordinates (rows)
    rows = {}
    for block in blocks:
        if block.get("type") != 0:  # Skip non-text
            continue
        bbox = block.get("bbox", [0, 0, 0, 0])
        y = round(bbox[1] / 5) * 5  # Group by Y position (5pt tolerance)

        if y not in rows:
            rows[y] = []
        rows[y].append(bbox)

    # If we have 3+ rows with 2+ columns each, likely a table
    row_list = sorted(rows.items())
    if len(row_list) >= 3:
        consecutive_table_rows = []
        for y, bboxes in row_list:
            if len(bboxes) >= 2:  # At least 2 columns
                consecutive_table_rows.append((y, bboxes))
            else:
                # Break in table
                if len(consecutive_table_rows) >= 3:
                    # Extract table bbox
                    all_bboxes = []
                    for _, bs in consecutive_table_rows:
                        all_bboxes.extend(bs)
                    x0 = min(b[0] for b in all_bboxes)
                    y0 = min(b[1] for b in all_bboxes)
                    x1 = max(b[2] for b in all_bboxes)
                    y1 = max(b[3] for b in all_bboxes)
                    tables.append((x0, y0, x1, y1))
                consecutive_table_rows = []

        # Check last group
        if len(consecutive_table_rows) >= 3:
            all_bboxes = []
            for _, bs in consecutive_table_rows:
                all_bboxes.extend(bs)
            x0 = min(b[0] for b in all_bboxes)
            y0 = min(b[1] for b in all_bboxes)
            x1 = max(b[2] for b in all_bboxes)
            y1 = max(b[3] for b in all_bboxes)
            tables.append((x0, y0, x1, y1))

    return tables


def _extract_table_data(
    page, table_bbox: tuple[float, float, float, float]
) -> list[list[str]]:
    """Extract table data from a detected table region.

    Returns 2D array of cell contents.
    """
    x0, y0, x1, y1 = table_bbox

    # Get all text blocks within table bbox
    blocks = (
        page.get_text("dict").get("blocks", [])
        if isinstance(page.get_text("dict"), dict)
        else []
    )

    table_blocks = []
    for block in blocks:
        if block.get("type") != 0:
            continue
        bbox = block.get("bbox", [0, 0, 0, 0])
        # Check if block is within table region
        if bbox[0] >= x0 and bbox[1] >= y0 and bbox[2] <= x1 and bbox[3] <= y1:
            table_blocks.append(block)

    # Group into rows by Y coordinate
    rows = {}
    for block in table_blocks:
        bbox = block.get("bbox", [0, 0, 0, 0])
        y = round(bbox[1] / 5) * 5

        # Extract text from block
        text = ""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                text += span.get("text", "") + " "
        text = text.strip()

        if y not in rows:
            rows[y] = []
        rows[y].append((bbox[0], text))  # (x_position, text)

    # Sort rows by Y, then sort cells in each row by X
    table_data = []
    for y in sorted(rows.keys()):
        cells = sorted(rows[y], key=lambda c: c[0])
        row = [cell[1] for cell in cells]
        table_data.append(row)

    return table_data


class PdfToDocxEngine(BaseEngine):
    """PDF to DOCX conversion engine using PyMuPDF + python-docx.

    Phase 2 - First version goals:
    - Preserve text order and basic structure
    - Make paragraphs editable
    - Export and embed images
    - Preserve page breaks
    - Not pixel-perfect, focus on editability
    """

    @property
    def name(self) -> str:
        return "PyMuPDF+python-docx"

    def can_convert(self, source_format: str, target_format: str) -> bool:
        """Check if this engine can handle the conversion."""
        return source_format == "pdf" and target_format == "docx"

    def convert(
        self,
        source_path: Path,
        source_format: str,
        target_format: str,
        output_path: Path,
    ) -> ConversionResult:
        """Convert PDF to DOCX with text and images."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            return ConversionResult(
                success=False,
                error_message="PyMuPDF not installed. Install with: pip install PyMuPDF",
            )

        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            return ConversionResult(
                success=False,
                error_message="python-docx not installed. Install with: pip install python-docx",
            )

        try:
            # Open PDF
            pdf_doc = fitz.open(source_path)
            word_doc = Document()

            page_count = len(pdf_doc)
            total_text_length = 0
            total_images = 0
            total_tables = 0
            image_temp_dir = output_path.parent / f".temp_images_{output_path.stem}"
            image_temp_dir.mkdir(exist_ok=True)

            # Process each page
            for page_num in range(page_count):
                page = pdf_doc[page_num]

                # Detect tables on this page
                table_regions = _detect_table_regions(page)

                # Extract text with detailed font information
                text_dict = page.get_text("dict")
                blocks = (
                    text_dict.get("blocks", []) if isinstance(text_dict, dict) else []
                )

                # Sort blocks by position
                sorted_blocks = sorted(
                    blocks,
                    key=lambda b: (
                        round(float(b.get("bbox", [0, 0, 0, 0])[1]) / 10) * 10,
                        float(b.get("bbox", [0, 0, 0, 0])[0]),
                    ),
                )

                # Track which blocks are inside tables
                blocks_in_tables = set()
                for block_idx, block in enumerate(sorted_blocks):
                    if block.get("type") != 0:
                        continue
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    for table_bbox in table_regions:
                        x0, y0, x1, y1 = table_bbox
                        if (
                            bbox[0] >= x0
                            and bbox[1] >= y0
                            and bbox[2] <= x1
                            and bbox[3] <= y1
                        ):
                            blocks_in_tables.add(block_idx)
                            break

                # Process tables first (by position)
                processed_table_regions = set()
                for block_idx, block in enumerate(sorted_blocks):
                    if (
                        block_idx not in blocks_in_tables
                        or block_idx in processed_table_regions
                    ):
                        continue

                    # Find which table this block belongs to
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    for table_bbox in table_regions:
                        x0, y0, x1, y1 = table_bbox
                        if (
                            bbox[0] >= x0
                            and bbox[1] >= y0
                            and bbox[2] <= x1
                            and bbox[3] <= y1
                        ):
                            # Extract and render table
                            table_data = _extract_table_data(page, table_bbox)
                            if table_data and len(table_data) > 0:
                                # Create Word table
                                from docx.oxml import OxmlElement
                                from docx.oxml.ns import qn

                                max_cols = max(len(row) for row in table_data)
                                table = word_doc.add_table(
                                    rows=len(table_data), cols=max_cols
                                )
                                table.style = "Table Grid"

                                for row_idx, row_data in enumerate(table_data):
                                    for col_idx, cell_text in enumerate(row_data):
                                        if col_idx < max_cols:
                                            table.rows[row_idx].cells[
                                                col_idx
                                            ].text = cell_text

                                total_tables += 1

                            # Mark all blocks in this table as processed
                            for other_idx, other_block in enumerate(sorted_blocks):
                                if other_idx in blocks_in_tables:
                                    other_bbox = other_block.get("bbox", [0, 0, 0, 0])
                                    if (
                                        other_bbox[0] >= x0
                                        and other_bbox[1] >= y0
                                        and other_bbox[2] <= x1
                                        and other_bbox[3] <= y1
                                    ):
                                        processed_table_regions.add(other_idx)
                            break

                # Process text blocks with font info (skip table blocks)
                for block_idx, block in enumerate(sorted_blocks):
                    if block.get("type") != 0:  # Skip non-text blocks
                        continue

                    # Skip blocks that are part of tables
                    if block_idx in blocks_in_tables:
                        continue

                    # Extract lines with font information
                    for line in block.get("lines", []):
                        line_text = ""
                        line_runs = []

                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue

                            line_text += text + " "
                            line_runs.append(
                                {
                                    "text": text,
                                    "font": span.get("font", ""),
                                    "size": span.get("size", 12),
                                    "flags": span.get(
                                        "flags", 0
                                    ),  # bit flags for bold/italic
                                    "color": span.get("color", 0),
                                }
                            )

                        line_text = line_text.strip()
                        if not line_text:
                            continue

                        total_text_length += len(line_text)

                        # Analyze line properties
                        avg_size = (
                            sum(r["size"] for r in line_runs) / len(line_runs)
                            if line_runs
                            else 12
                        )
                        has_bold = any(
                            r["flags"] & 2**4 for r in line_runs
                        )  # bit 4 = bold

                        # Heading detection: large font or short bold text
                        is_heading = (avg_size > 14) or (
                            len(line_text) < 80 and has_bold
                        )

                        # Add paragraph with formatting
                        if is_heading:
                            para = word_doc.add_paragraph(line_text)
                            para.style = "Heading 2"
                        else:
                            para = word_doc.add_paragraph()

                            # Apply run-level formatting
                            for run_info in line_runs:
                                run = para.add_run(run_info["text"] + " ")
                                run.font.size = Pt(run_info["size"])
                                if run_info["flags"] & 2**4:  # bold
                                    run.bold = True
                                if run_info["flags"] & 2**1:  # italic
                                    run.italic = True

                # Extract and embed images with bbox-based sizing
                image_list = page.get_images(full=True)
                for img_index, img_info in enumerate(image_list):
                    try:
                        xref = img_info[0]
                        base_image = pdf_doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        # Get image bounding box from page
                        image_rects = page.get_image_rects(xref)
                        img_width_inches = 5.0  # Default width

                        if image_rects:
                            # Use first rect (images can appear multiple times)
                            rect = image_rects[0]
                            # rect is (x0, y0, x1, y1) in points
                            bbox_width = rect[2] - rect[0]  # in points
                            bbox_height = rect[3] - rect[1]  # in points

                            # Convert points to inches (72 points = 1 inch)
                            img_width_inches = bbox_width / 72
                            img_height_inches = bbox_height / 72

                            # Clamp to reasonable page margins (max 6.5 inches for standard Letter page)
                            img_width_inches = min(img_width_inches, 6.5)
                            img_height_inches = min(img_height_inches, 9.0)

                        # Save temp image
                        image_filename = (
                            f"page{page_num + 1}_img{img_index + 1}.{image_ext}"
                        )
                        image_path = image_temp_dir / image_filename
                        image_path.write_bytes(image_bytes)

                        # Add to Word document with bbox-based size
                        try:
                            if image_rects:
                                word_doc.add_picture(
                                    str(image_path), width=Inches(img_width_inches)
                                )
                            else:
                                word_doc.add_picture(str(image_path), width=Inches(5.0))
                            total_images += 1
                        except Exception as e:
                            logger.warning(f"Failed to add image {image_filename}: {e}")

                    except Exception as e:
                        logger.warning(
                            f"Failed to extract image {img_index} from page {page_num}: {e}"
                        )

                # Add page break (except for last page)
                if page_num < page_count - 1:
                    word_doc.add_page_break()

            # Save Word document
            word_doc.save(str(output_path))

            # Close PDF before building quality report
            pdf_doc.close()

            # Cleanup temp images
            try:
                for img_file in image_temp_dir.glob("*"):
                    img_file.unlink()
                image_temp_dir.rmdir()
            except Exception as e:
                logger.warning(f"Failed to cleanup temp images: {e}")

            # Build quality report
            quality_report = self._build_quality_report(
                page_count=page_count,
                text_length=total_text_length,
                image_count=total_images,
                table_count=total_tables,
                output_path=output_path,
            )

            logger.info(
                f"PDF → DOCX conversion complete: {page_count} pages, "
                f"{total_text_length} chars, {total_images} images, {total_tables} tables"
            )

            return ConversionResult(
                success=True,
                output_path=output_path,
                metadata={
                    "engine": self.name,
                    "page_count": page_count,
                    "text_length": total_text_length,
                    "image_count": total_images,
                    "table_count": total_tables,
                    "quality_report": quality_report,
                },
            )

        except Exception as e:
            logger.error(f"PDF → DOCX conversion failed: {e}", exc_info=True)
            return ConversionResult(
                success=False,
                error_message=f"PDF → DOCX conversion failed: {str(e)}",
            )

    def _build_quality_report(
        self,
        page_count: int,
        text_length: int,
        image_count: int,
        table_count: int,
        output_path: Path,
    ) -> dict[str, Any]:
        """Build quality report for conversion result."""
        issues = []
        quality_level = "high"

        # Check if output file exists and has size
        if not output_path.exists():
            issues.append("Output file not created")
            quality_level = "low"
        else:
            file_size = output_path.stat().st_size
            if file_size < 1000:  # Less than 1KB is suspicious
                issues.append("Output file suspiciously small")
                quality_level = "low"

            if text_length < 100:
                issues.append("Very little text extracted")
                quality_level = "medium"

        return {
            "quality_level": quality_level,
            "text_length": text_length,
            "image_count": image_count,
            "table_count": table_count,
            "output_file_size": output_path.stat().st_size
            if output_path.exists()
            else 0,
            "issues": issues,
        }
