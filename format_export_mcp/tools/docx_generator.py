from __future__ import annotations

from pathlib import Path

from .image_sources import load_image_assets
from .markdown_blocks import parse_markdown_blocks, parse_markdown_inlines


def generate_docx(title: str, content: str, output_path: Path, images: list[str] | None = None) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt

    def _set_east_asia_font(run, font_name: str) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _add_inline_runs(paragraph, text: str, font_name: str = "Microsoft YaHei", font_size: int = 11) -> None:
        for span in parse_markdown_inlines(text):
            run = paragraph.add_run(span.text)
            if "bold" in span.styles:
                run.bold = True
            if "italic" in span.styles:
                run.italic = True
            if "underline" in span.styles:
                run.underline = True
            if "strike" in span.styles:
                run.font.strike = True
            if "code" in span.styles:
                _set_east_asia_font(run, "Consolas")
            else:
                _set_east_asia_font(run, font_name)
            run.font.size = Pt(font_size)

    def _add_image(image_ref: str) -> None:
        image_asset = load_image_assets([image_ref])[0]
        document.add_picture(image_asset.open_bytes(), width=Mm(160))

    document = Document()
    document.core_properties.title = title or "Export"

    heading = document.add_heading(level=1)
    _add_inline_runs(heading, title or "Export", font_size=18)

    for block in parse_markdown_blocks(content or "") or [None]:
        if block is None:
            paragraph = document.add_paragraph()
            _add_inline_runs(paragraph, "", font_size=11)
            continue

        if block.kind == "heading":
            paragraph = document.add_heading(level=min(block.level, 4))
            _add_inline_runs(paragraph, block.text, font_size=max(12, 20 - (block.level * 2)))
            continue

        if block.kind == "bullet_item":
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_runs(paragraph, block.text, font_size=11)
            continue

        if block.kind == "ordered_item":
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_runs(paragraph, block.text, font_size=11)
            continue

        if block.kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(block.text)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            continue

        if block.kind == "image" and block.image_src:
            _add_image(block.image_src)
            continue

        if block.kind == "table" and block.rows:
            table = document.add_table(rows=len(block.rows), cols=max(len(row) for row in block.rows))
            table.style = "Table Grid"
            for row_index, row in enumerate(block.rows):
                for col_index, cell_text in enumerate(row):
                    cell = table.rows[row_index].cells[col_index]
                    paragraph = cell.paragraphs[0]
                    _add_inline_runs(paragraph, cell_text, font_size=10)
                    if row_index == 0:
                        for run in paragraph.runs:
                            run.bold = True
            continue

        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, block.text, font_size=11)

    for image_asset in load_image_assets(list(images or [])):
        document.add_picture(image_asset.open_bytes(), width=Mm(160))

    document.save(output_path)
