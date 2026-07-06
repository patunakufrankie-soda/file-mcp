from __future__ import annotations

from pathlib import Path

from ...utils.image_sources import load_image_assets
from ...utils.markdown_blocks import (
    parse_markdown_blocks,
    parse_markdown_inlines,
    parse_plain_text_blocks,
    render_markdown_inlines_as_text,
)


def generate_docx(
    title: str,
    content: str,
    output_path: Path,
    images: list[str] | None = None,
    input_format: str = "md",
) -> None:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Mm, Pt, RGBColor

    def _set_east_asia_font(run, font_name: str) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _add_inline_runs(
        paragraph, text: str, font_name: str = "Microsoft YaHei", font_size: int = 11
    ) -> None:
        for span in parse_markdown_inlines(text):
            run = paragraph.add_run(span.text)
            if "bold" in span.styles:
                run.bold = True
            if "italic" in span.styles:
                run.italic = True
            if "underline" in span.styles:
                run.underline = True
            if "strike" in span.styles:
                run.font.strike = True
            if "code" in span.styles:
                _set_east_asia_font(run, "Consolas")
            else:
                _set_east_asia_font(run, font_name)
            run.font.size = Pt(font_size)
            if span.href:
                relationship_id = paragraph.part.relate_to(
                    span.href,
                    RELATIONSHIP_TYPE.HYPERLINK,
                    is_external=True,
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), relationship_id)
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
                hyperlink.append(run._element)
                paragraph._p.append(hyperlink)

    def _add_image(image_ref: str) -> None:
        image_asset = load_image_assets([image_ref])[0]
        document.add_picture(image_asset.open_bytes(), width=Mm(160))

    document = Document()
    document.core_properties.title = title or "Export"

    if "Intense Quote" not in document.styles:
        quote_style = document.styles.add_style(
            "Intense Quote",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        quote_style.font.italic = True

    heading = document.add_heading(level=1)
    _add_inline_runs(heading, title or "Export", font_size=18)

    blocks = (
        parse_plain_text_blocks(content or "")
        if input_format == "txt"
        else parse_markdown_blocks(content or "")
    )
    if (
        blocks
        and blocks[0].kind == "heading"
        and render_markdown_inlines_as_text(blocks[0].text).strip()
        == (title or "Export").strip()
    ):
        blocks = blocks[1:]

    for block in blocks or [None]:
        if block is None:
            paragraph = document.add_paragraph()
            _add_inline_runs(paragraph, "", font_size=11)
            continue

        if block.kind == "heading":
            paragraph = document.add_heading(level=min(block.level, 6))
            _add_inline_runs(
                paragraph, block.text, font_size=max(12, 20 - (block.level * 2))
            )
            continue

        if block.kind == "bullet_item":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Pt(18 * (block.depth + 1))
            paragraph.paragraph_format.first_line_indent = Pt(-9)
            _add_inline_runs(paragraph, block.text, font_size=11)
            continue

        if block.kind == "ordered_item":
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Pt(18 * (block.depth + 1))
            paragraph.paragraph_format.first_line_indent = Pt(-9)
            _add_inline_runs(paragraph, block.text, font_size=11)
            continue

        if block.kind == "blockquote":
            paragraph = document.add_paragraph(style="Intense Quote")
            _add_inline_runs(paragraph, block.text, font_size=11)
            continue

        if block.kind == "horizontal_rule":
            paragraph = document.add_paragraph()
            paragraph_properties = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "9CA3AF")
            borders.append(bottom)
            paragraph_properties.append(borders)
            continue

        if block.kind == "code":
            paragraph = document.add_paragraph()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F3F4F6")
            paragraph._p.get_or_add_pPr().append(shading)
            run = paragraph.add_run(block.text)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
            continue

        if block.kind == "image" and block.image_src:
            _add_image(block.image_src)
            continue

        if block.kind == "table" and block.rows:
            table = document.add_table(
                rows=len(block.rows), cols=max(len(row) for row in block.rows)
            )
            table.style = "Table Grid"
            for row_index, row in enumerate(block.rows):
                for col_index, cell_text in enumerate(row):
                    cell = table.rows[row_index].cells[col_index]
                    paragraph = cell.paragraphs[0]
                    _add_inline_runs(paragraph, cell_text, font_size=10)
                    if row_index == 0:
                        for run in paragraph.runs:
                            run.bold = True
            header_properties = table.rows[0]._tr.get_or_add_trPr()
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            header_properties.append(repeat_header)
            continue

        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, block.text, font_size=11)

    for image_asset in load_image_assets(list(images or [])):
        document.add_picture(image_asset.open_bytes(), width=Mm(160))

    document.save(output_path)
