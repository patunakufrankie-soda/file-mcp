#!/usr/bin/env python3
"""Test enhanced conversion features with logging and quality assessment."""

from pathlib import Path
from format_export_mcp.conversion.services.conversion_service import ConversionService
from format_export_mcp.conversion.services.format_detector import FormatDetector
import tempfile
import logging

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")


def test_pdf_analysis():
    """Test enhanced PDF analysis."""
    print("\n=== Test 1: Enhanced PDF Analysis ===")

    # Create a test PDF with text
    test_pdf = Path("test_files/sample.pdf")
    if not test_pdf.exists():
        print(f"⚠ Test PDF not found: {test_pdf}")
        print("Creating a simple test PDF...")

        # Create test directory
        test_pdf.parent.mkdir(exist_ok=True)

        # Create a simple PDF using reportlab if available
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter

            c = canvas.Canvas(str(test_pdf), pagesize=letter)
            c.drawString(100, 750, "Test PDF Document")
            c.drawString(100, 700, "This is a simple text-based PDF for testing.")
            c.drawString(100, 650, "It contains multiple text blocks.")
            c.showPage()
            c.save()
            print(f"✓ Created test PDF: {test_pdf}")
        except ImportError:
            print("⚠ reportlab not available, skipping PDF creation")
            return

    # Analyze PDF
    detector = FormatDetector()
    format_str = detector.detect_format(test_pdf)
    features = detector.analyze_features(test_pdf, format_str)

    print(f"\nFormat: {format_str}")
    print(f"Page count: {features.page_count}")
    print(f"Has text layer: {features.has_text_layer}")
    print(f"Is scanned: {features.is_scanned}")
    print(f"Image count: {features.image_count}")
    print(f"Text block count: {features.text_block_count}")
    print(f"Table count: {features.table_count}")
    print(f"Multi-column: {features.has_multicolumn}")
    print(f"Complexity: {features.estimated_complexity}")
    print(f"Analysis method: {features.analysis_method}")
    print(f"Confidence: {features.analysis_confidence}")
    print(f"Text coverage: {features.text_coverage_ratio:.2%}")


def test_libreoffice_conversion():
    """Test LibreOffice-based conversions with logging."""
    print("\n=== Test 2: LibreOffice Conversion with Logging ===")

    # Create a test DOCX
    test_docx = Path("test_files/sample.docx")
    if not test_docx.exists():
        print(f"Creating test DOCX: {test_docx}")
        test_docx.parent.mkdir(exist_ok=True)

        try:
            from docx import Document

            doc = Document()
            doc.add_heading("Test Document", 0)
            doc.add_paragraph("This is a test paragraph.")
            doc.add_paragraph("Another paragraph with more content.")
            doc.save(str(test_docx))
            print(f"✓ Created test DOCX: {test_docx}")
        except ImportError:
            print("⚠ python-docx not available, skipping DOCX creation")
            return

    # Convert DOCX to PDF with logging
    with tempfile.TemporaryDirectory() as tmpdir:
        log_file = Path(tmpdir) / "conversion.log"
        service = ConversionService(log_file=log_file)

        output_pdf = Path(tmpdir) / "output.pdf"
        result = service.convert(test_docx, "pdf", output_pdf)

        print(f"\nConversion result:")
        print(f"  Success: {result.success}")
        if result.success and result.output_path:
            print(f"  Output: {result.output_path}")
            print(f"  Output size: {result.output_path.stat().st_size} bytes")
        else:
            print(f"  Error: {result.error_message}")

        # Show log
        if log_file.exists():
            print(f"\nConversion log:")
            print(log_file.read_text())


def test_pdf_to_text_conversion():
    """Test PDF to text conversion with quality assessment."""
    print("\n=== Test 3: PDF → Text with Quality Assessment ===")

    test_pdf = Path("test_files/sample.pdf")
    if not test_pdf.exists():
        print(f"⚠ Test PDF not found, skipping")
        return

    with tempfile.TemporaryDirectory() as tmpdir:
        log_file = Path(tmpdir) / "conversion.log"
        service = ConversionService(log_file=log_file)

        # Convert to TXT
        output_txt = Path(tmpdir) / "output.txt"
        result = service.convert(test_pdf, "txt", output_txt)

        print(f"\nPDF → TXT:")
        print(f"  Success: {result.success}")
        if result.success and result.output_path:
            content = result.output_path.read_text(encoding="utf-8")
            print(f"  Text length: {len(content)} chars")
            print(f"  Preview: {content[:200]}...")

        # Convert to MD
        output_md = Path(tmpdir) / "output.md"
        result = service.convert(test_pdf, "md", output_md)

        print(f"\nPDF → MD:")
        print(f"  Success: {result.success}")
        if result.success and result.output_path:
            content = result.output_path.read_text(encoding="utf-8")
            print(f"  Markdown length: {len(content)} chars")
            print(f"  Preview: {content[:200]}...")


def test_edge_cases():
    """Test edge cases: empty file, non-existent file, unsupported format."""
    print("\n=== Test 4: Edge Cases ===")

    service = ConversionService()

    # Non-existent file
    print("\n1. Non-existent file:")
    result = service.convert(Path("nonexistent.pdf"), "txt")
    print(f"   Success: {result.success}")
    print(f"   Error: {result.error_message}")

    # Unsupported conversion
    print("\n2. Unsupported conversion (jpg → pdf):")
    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
        test_jpg = Path(f.name)
        f.write(b"fake image data")

    result = service.convert(test_jpg, "pdf")
    print(f"   Success: {result.success}")
    print(f"   Error: {result.error_message}")
    test_jpg.unlink()


if __name__ == "__main__":
    print("Testing Enhanced Conversion System")
    print("=" * 50)

    test_pdf_analysis()
    test_libreoffice_conversion()
    test_pdf_to_text_conversion()
    test_edge_cases()

    print("\n" + "=" * 50)
    print("✓ All tests completed")
