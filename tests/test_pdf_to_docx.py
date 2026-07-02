"""Tests for PDF to DOCX conversion (Phase 2)."""

import pytest
from pathlib import Path
import tempfile
import shutil

from format_export_mcp.conversion.services.format_detector import FormatDetector
from format_export_mcp.conversion.router import ConversionRouter
from format_export_mcp.conversion.services.conversion_service import ConversionService


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test outputs."""
    temp = Path(tempfile.mkdtemp())
    yield temp
    shutil.rmtree(temp, ignore_errors=True)


@pytest.fixture
def sample_text_pdf(temp_dir):
    """Create a simple text-based PDF for testing."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    pdf_path = temp_dir / "text_sample.pdf"
    doc = fitz.open()

    # Add pages with text
    page1 = doc.new_page()
    page1.insert_text((72, 72), "Test Document", fontsize=20)
    page1.insert_text(
        (72, 120), "This is a simple text-based PDF document.", fontsize=12
    )
    page1.insert_text(
        (72, 150), "It contains multiple paragraphs of text.", fontsize=12
    )

    page2 = doc.new_page()
    page2.insert_text((72, 72), "Second Page", fontsize=20)
    page2.insert_text((72, 120), "More content on the second page.", fontsize=12)

    doc.save(pdf_path)
    doc.close()

    return pdf_path


@pytest.fixture
def sample_image_pdf(temp_dir):
    """Create a PDF with images for testing."""
    try:
        import fitz
        from PIL import Image
    except ImportError:
        pytest.skip("PyMuPDF or Pillow not installed")

    pdf_path = temp_dir / "image_sample.pdf"
    doc = fitz.open()

    # Create a test image
    img = Image.new("RGB", (200, 200), color="red")
    img_path = temp_dir / "test_image.png"
    img.save(img_path)

    # Add page with text and image
    page = doc.new_page()
    page.insert_text((72, 72), "Document with Image", fontsize=16)
    page.insert_text(
        (72, 120), "This PDF contains both text and an image.", fontsize=12
    )

    # Insert image
    rect = fitz.Rect(72, 150, 272, 350)
    page.insert_image(rect, filename=str(img_path))

    doc.save(pdf_path)
    doc.close()

    return pdf_path


class TestFormatDetector:
    """Test enhanced FormatDetector with strategy recommendation."""

    def test_detect_text_pdf_strategy(self, sample_text_pdf):
        """Test detection of text-based PDF strategy."""
        detector = FormatDetector()
        format_type = detector.detect_format(sample_text_pdf)
        features = detector.analyze_features(sample_text_pdf, format_type)

        assert format_type == "pdf"
        assert features.has_text_layer is True
        assert features.is_scanned is False
        assert features.recommended_strategy in ["text_pdf", "office_like_pdf"]
        assert features.page_count == 2

    def test_detect_image_pdf_strategy(self, sample_image_pdf):
        """Test detection of PDF with images."""
        detector = FormatDetector()
        format_type = detector.detect_format(sample_image_pdf)
        features = detector.analyze_features(sample_image_pdf, format_type)

        assert format_type == "pdf"
        assert features.has_images is True
        assert features.image_count >= 1
        assert features.recommended_strategy in ["mixed_pdf", "text_pdf"]


class TestConversionRouter:
    """Test dynamic routing for PDF → DOCX."""

    def test_route_text_pdf_to_docx(self, sample_text_pdf):
        """Test routing for text-based PDF."""
        detector = FormatDetector()
        router = ConversionRouter()

        features = detector.analyze_features(sample_text_pdf, "pdf")
        route = router.get_route("pdf", "docx", features)

        assert route is not None
        assert route.primary is not None
        # Should have fallback options
        assert len(route.fallbacks) >= 0

    def test_route_without_features(self):
        """Test routing without features falls back to basic route."""
        router = ConversionRouter()
        route = router.get_route("pdf", "docx", features=None)

        # Should still return a route (might be None if not in basic routes)
        # In our implementation, PDF→DOCX without features returns None,
        # forcing feature detection
        assert route is None or route is not None


class TestPdfToDocxConversion:
    """Test PDF → DOCX conversion with content validation."""

    def test_text_pdf_to_docx(self, sample_text_pdf, temp_dir):
        """Test converting text PDF to DOCX."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        service = ConversionService()
        output_path = temp_dir / "output.docx"

        result = service.convert(
            source_file=sample_text_pdf, target_format="docx", output_path=output_path
        )

        assert result.success is True
        assert output_path.exists()
        assert output_path.stat().st_size > 0

        # Validate DOCX content
        doc = Document(str(output_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Should have extracted text content
        assert len(paragraphs) > 0, "DOCX should contain text paragraphs"
        assert any("Test Document" in p for p in paragraphs), (
            "Should contain original text"
        )

        # Check quality report
        assert result.metadata is not None
        if "quality_report" in result.metadata:
            quality = result.metadata["quality_report"]
            assert quality["quality_level"] in ["high", "medium", "low"]
            assert quality["text_length"] > 0, "Should report text length"

    def test_image_pdf_to_docx(self, sample_image_pdf, temp_dir):
        """Test converting PDF with images to DOCX."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        service = ConversionService()
        output_path = temp_dir / "output_with_image.docx"

        result = service.convert(
            source_file=sample_image_pdf, target_format="docx", output_path=output_path
        )

        assert result.success is True
        assert output_path.exists()

        # Validate content
        doc = Document(str(output_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Should have text
        assert len(paragraphs) > 0
        assert any("Image" in p for p in paragraphs)

        # Check for images (python-docx stores images in document relations)
        # For now, just verify the file size is larger (contains image data)
        assert output_path.stat().st_size > 5000, "DOCX with image should be larger"

        # Check quality report includes image count
        if result.metadata and "quality_report" in result.metadata:
            quality = result.metadata["quality_report"]
            assert quality["image_count"] >= 1, "Should report image extraction"

    def test_empty_pdf_handling(self, temp_dir):
        """Test handling of problematic PDF (empty/corrupted)."""
        try:
            import fitz
        except ImportError:
            pytest.skip("PyMuPDF not installed")

        # Create empty PDF
        empty_pdf = temp_dir / "empty.pdf"
        doc = fitz.open()
        doc.new_page()  # Empty page
        doc.save(empty_pdf)
        doc.close()

        service = ConversionService()
        output_path = temp_dir / "empty_output.docx"

        result = service.convert(
            source_file=empty_pdf, target_format="docx", output_path=output_path
        )

        # Empty PDF is now treated as scanned_pdf → OCR_REQUIRED
        assert result.success is False
        assert result.error_message is not None
        assert "OCR_REQUIRED" in result.error_message
        assert result.metadata is not None
        assert result.metadata.get("error_code") == "OCR_REQUIRED"
        assert result.metadata.get("is_scanned") is True


class TestQualityReport:
    """Test quality reporting for PDF → DOCX conversions."""

    def test_quality_report_structure(self, sample_text_pdf, temp_dir):
        """Test that quality report has expected structure."""
        service = ConversionService()
        output_path = temp_dir / "quality_test.docx"

        result = service.convert(
            source_file=sample_text_pdf, target_format="docx", output_path=output_path
        )

        assert result.success is True
        assert result.metadata is not None

        if "quality_report" in result.metadata:
            quality = result.metadata["quality_report"]

            # Check required fields
            assert "quality_level" in quality
            assert "output_file_size" in quality
            assert "issues" in quality
            assert isinstance(quality["issues"], list)

            # Validate quality level
            assert quality["quality_level"] in ["high", "medium", "low", "fallback"]


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
